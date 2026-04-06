import streamlit as st
from groq import Groq
from docx import Document
from io import BytesIO
import re
import pandas as pd
from datetime import datetime

# --- 1. BRANDING & STYLE ---
st.set_page_config(page_title="VIKIDYL AI Pro", page_icon="🎓", layout="wide")

st.markdown("""
    <style>
    .stButton>button { width: 100%; border-radius: 25px; background-color: #1E88E5; color: white; height: 3.5em; font-weight: bold; }
    .footer { position: fixed; left: 0; bottom: 0; width: 100%; background-color: #f1f1f1; text-align: center; padding: 10px; font-size: 14px; border-top: 1px solid #e0e0e0; z-index: 99; }
    </style>
    """, unsafe_allow_html=True)

# --- 2. LIVE SECURITY LAYER ---
SHEET_URL = "https://docs.google.com/spreadsheets/d/1Q3UQhyOt8YS800kJ2aMw3bbsVpji-JUy_jU1shoeHYs/edit#gid=0"

def get_live_codes():
    try:
        base_url = SHEET_URL.split('/edit')[0]
        csv_url = f"{base_url}/export?format=csv"
        df = pd.read_csv(csv_url)
        df.columns = [c.strip().lower() for c in df.columns]
        return dict(zip(df['code'].astype(str), df['expiry'].astype(str)))
    except:
        return {"VIK-ADMIN-2026": "2027-01-01"}

def check_access(user_input):
    valid_codes = get_live_codes()
    if user_input in valid_codes:
        try:
            expiry_str = valid_codes[user_input].strip()
            expiry_date = datetime.strptime(expiry_str, "%Y-%m-%d")
            if datetime.now() <= expiry_date:
                return True, expiry_date.strftime("%d %B %Y")
            else:
                return False, "Expired"
        except:
            return False, "Date Format Error"
    return False, "Invalid Code"

if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

if not st.session_state['authenticated']:
    st.title("🔐 VIKIDYL AI - Secure Access")
    user_code = st.text_input("Enter Teacher Access Code", type="password")
    if st.button("Unlock System"):
        is_valid, status = check_access(user_code)
        if is_valid:
            st.session_state['authenticated'] = True
            st.session_state['expiry'] = status
            st.rerun()
        else:
            st.error(f"Access Denied: {status}")
    st.stop()

# --- 3. THE "READY-TO-PRINT" MATH ENGINE ---
def clean_math_output(text):
    """Forcefully removes LaTeX and converts to readable teacher symbols"""
    # 1. Handle Fractions: \frac{a}{b} -> (a / b)
    text = re.sub(r'\\frac\{(.+?)\}\{(.+?)\}', r'(\1 / \2)', text)
    # 2. Handle Roots: \sqrt{x} -> √(x)
    text = re.sub(r'\\sqrt\{(.+?)\}', r'√(\1)', text)
    # 3. Handle Common LaTeX Symbols
    replacements = {
        r'\\times': '×',
        r'\\div': '÷',
        r'\\pm': '±',
        r'\\geq': '≥',
        r'\\leq': '≤',
        r'\\neq': '≠',
        r'\\approx': '≈',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\theta': 'θ',
        r'\\pi': 'π',
        r'\\sum': 'Σ',
        r'\\infty': '∞',
        r'\^2': '²',
        r'\^3': '³',
        r'\$': '', # Remove dollar signs
        r'\\': ''  # Catch-all for remaining backslashes
    }
    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text)
    return text

def create_docx(text, title):
    doc = Document()
    doc.add_heading(f'VIKIDYL AI Professional - {title}', 0)
    # Clean math before putting it into the Word Doc
    clean_text = clean_math_output(text)
    for line in clean_text.split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. CURRICULUM DATA ---
CURRICULUM_DATA = {
    "Nursery": {"Classes": ["Pre-Nursery", "Nursery 1", "Nursery 2"], "Subjects": ["Letter Work", "Number Work", "Health Habits", "Social Norms", "Rhymes", "Science Experience"]},
    "Primary": {"Classes": ["Basic 1", "Basic 2", "Basic 3", "Basic 4", "Basic 5", "Basic 6"], "Subjects": ["Mathematics", "English Studies", "Basic Science & Tech", "Social Studies", "Nigerian History", "P.H.E", "Agriculture", "Home Economics"]},
    "Junior Secondary": {"Classes": ["JSS 1", "JSS 2", "JSS 3"], "Subjects": ["Mathematics", "English", "Basic Science", "Basic Technology", "Business Studies", "Civic Education", "Social Studies"]},
    "Senior Secondary": {"Classes": ["SSS 1", "SSS 2", "SSS 3"], "Subjects": ["Mathematics", "English Language", "Biology", "Chemistry", "Physics", "Further Maths", "Economics", "Government", "Food and Nutrition", "Financial Accounting"]}
}

# --- 5. MAIN APP ---
try:
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
except:
    st.error("🔑 API Key Error: Check your Streamlit Secrets.")
    st.stop()

st.sidebar.success(f"Access Active until {st.session_state['expiry']}")
if st.sidebar.button("Logout"):
    st.session_state['authenticated'] = False
    st.rerun()

st.title("🎓 VIKIDYL AI Professional")
tabs = st.tabs(["📝 Quick Tools", "📊 Assessment Builder", "📅 NERDC Scheme & Notes"])

# TAB 1 & 2 (Cleaned Math included here too)
with tabs[0]:
    st.subheader("Fast Lesson Script")
    lvl_q = st.selectbox("Level Group", list(CURRICULUM_DATA.keys()), key="lvl_q")
    cls_q = st.selectbox("Class", CURRICULUM_DATA[lvl_q]["Classes"], key="cls_q")
    sub_q = st.selectbox("Subject", CURRICULUM_DATA[lvl_q]["Subjects"], key="sub_q")
    top_q = st.text_input("Topic", key="top_q")
    if st.button("Generate Script"):
        with st.spinner("Drafting..."):
            chat = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "user", "content": f"Lesson script for {cls_q} {sub_q}: {top_q}. Use plain symbols only."}])
            st.session_state['out'] = clean_math_output(chat.choices[0].message.content)

with tabs[1]:
    st.subheader("Assessment Builder")
    c1, c2 = st.columns(2)
    with c1:
        lvl_a = st.selectbox("Level Group", list(CURRICULUM_DATA.keys()), key="lvl_a")
        cls_a = st.selectbox("Class", CURRICULUM_DATA[lvl_a]["Classes"], key="cls_a")
        sub_a = st.selectbox("Subject", CURRICULUM_DATA[lvl_a]["Subjects"], key="sub_a")
    with c2:
        diff_a = st.selectbox("Tone", ["Standard", "Moderate", "Fun"], key="diff_a")
        top_a = st.text_area("Topics", key="top_a")
    if st.button("Generate Assessment"):
        with st.spinner("Creating questions..."):
            chat = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "user", "content": f"Create {diff_a} exam for {cls_a} {sub_a} on: {top_a}. Use plain text math symbols."}])
            st.session_state['out'] = clean_math_output(chat.choices[0].message.content)

# --- TAB 3: NERDC (FULL 18-POINT + NO-CODE MATH) ---
with tabs[2]:
    st.subheader("NERDC Serialized Planner")
    col1, col2 = st.columns(2)
    with col1:
        lvl_f = st.selectbox("Level Group", list(CURRICULUM_DATA.keys()), key="lvl_f")
        cls_f = st.selectbox("Class", CURRICULUM_DATA[lvl_f]["Classes"], key="cls_f")
        sub_f = st.selectbox("Subject", CURRICULUM_DATA[lvl_f]["Subjects"], key="sub_f")
        trm_f = st.selectbox("Term", ["1st Term", "2nd Term", "3rd Term", "Full Year (All 3 Terms)"], key="trm_f")
    with col2:
        mode_f = st.radio("Action", ["Serialized Scheme (Week 1-12)", "Detailed Weekly Lesson Note"])
        manual_ref = st.text_area("Manual Reference (Optional Override)", height=70)
    
    if st.button("🚀 Generate NERDC Material"):
        age_guardrail = "Toddler-friendly play words only." if cls_f == "Pre-Nursery" else "Academic NERDC standards."

        if mode_f == "Serialized Scheme (Week 1-12)":
            p = f"Scheme for {cls_f} {sub_f} {trm_f}. 12 individual weeks. No grouping. {age_guardrail} {manual_ref}"
        else:
            p = f"""Professional 18-Point Lesson Note for {cls_f} {sub_f} ({trm_f}).
            Topic: {manual_ref if manual_ref else 'Appropriate NERDC Topic'}.
            {age_guardrail}
            STRICTLY USE PLAIN SYMBOLS (x, /, ^2, √). NEVER USE BACKSLASHES OR LATEX CODE.
            FOLLOW 18-POINT STRUCTURE: 1-6 Admin/Objectives, 7-12 Steps, 13 Worked Examples, 14-15 Exercises, 16-17 Eval/Conclusion, 18 Full Student Note."""

        with st.spinner("Generating Final Document..."):
            chat = client.chat.completions.create(
                model="llama-3.3-70b-versatile", 
                messages=[
                    {"role": "system", "content": "You are a Nigerian NERDC Specialist. You are FORBIDDEN from using LaTeX math code. You must use only symbols available on a standard keyboard or Unicode (like ×, ÷, ², √)."},
                    {"role": "user", "content": p}
                ]
            )
            # Final scrub for both screen and download
            st.session_state['out'] = clean_math_output(chat.choices[0].message.content)

# --- OUTPUT DISPLAY & DOWNLOAD ---
if 'out' in st.session_state:
    st.markdown("---")
